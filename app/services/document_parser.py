"""
文档解析模块
支持 PDF 和 Word 文档的文本提取与清洗
"""

import re
from typing import Optional
from fastapi import UploadFile, HTTPException
import fitz  # PyMuPDF
from docx import Document
from app.core.logger import app_logger


class DocumentParser:
    """
    文档解析器
    
    支持 PDF 和 Word 格式的文档解析，提取纯文本内容并进行清洗。
    
    Attributes:
        SUPPORTED_EXTENSIONS: 支持的文件扩展名集合
    """
    
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}
    
    def __init__(self):
        """初始化文档解析器"""
        app_logger.info("📄 DocumentParser 初始化完成")
    
    async def parse_file(self, file: UploadFile) -> str:
        """
        解析上传的文件，提取纯文本内容
        
        Args:
            file: FastAPI 上传的文件对象
            
        Returns:
            str: 清洗后的文本内容
            
        Raises:
            HTTPException: 文件格式不支持或解析失败
        """
        # 获取文件扩展名
        filename = file.filename or ""
        file_ext = self._get_file_extension(filename)
        
        # 验证文件格式
        if file_ext not in self.SUPPORTED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件格式: {file_ext}。支持的格式: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )
        
        app_logger.info(f"📂 开始解析文件: {filename} (格式: {file_ext})")
        
        try:
            # 读取文件内容
            content = await file.read()

            # 拒绝 0 字节文件
            if len(content) == 0:
                raise HTTPException(
                    status_code=400,
                    detail="文件内容为空（0 字节），请上传有效的合同文件"
                )

            # 根据文件类型选择解析方法
            if file_ext == ".pdf":
                text = self._parse_pdf(content)
            elif file_ext in {".docx", ".doc"}:
                text = self._parse_docx(content)
            else:
                raise HTTPException(status_code=400, detail=f"未实现的文件格式: {file_ext}")
            
            # 清洗文本
            cleaned_text = self._clean_text(text)
            
            app_logger.info(f"✅ 文件解析成功: {filename}, 提取字符数: {len(cleaned_text)}")
            return cleaned_text
            
        except HTTPException:
            raise
        except Exception as e:
            app_logger.error(f"❌ 文件解析失败: {filename}, 错误: {str(e)}")
            raise HTTPException(status_code=500, detail=f"文件解析失败: {str(e)}")
    
    def _get_file_extension(self, filename: str) -> str:
        """
        获取文件扩展名（小写）
        
        Args:
            filename: 文件名
            
        Returns:
            str: 小写的文件扩展名，如 ".pdf"
        """
        if not filename:
            return ""
        
        # 提取扩展名并转为小写
        parts = filename.rsplit(".", 1)
        if len(parts) > 1:
            return f".{parts[1].lower()}"
        return ""
    
    def _parse_pdf(self, content: bytes) -> str:
        """
        解析 PDF 文件内容
        
        使用 PyMuPDF (fitz) 提取 PDF 中的文本内容。
        
        Args:
            content: PDF 文件的二进制内容
            
        Returns:
            str: 提取的文本内容
            
        Raises:
            Exception: 当 PDF 解析失败时抛出异常
        """
        text_parts = []
        
        try:
            # 使用 PyMuPDF 打开 PDF
            with fitz.open(stream=content, filetype="pdf") as pdf_doc:
                # 遍历每一页
                for page_num in range(len(pdf_doc)):
                    page = pdf_doc[page_num]
                    # 提取页面文本
                    page_text = page.get_text()
                    if page_text.strip():
                        text_parts.append(page_text)
            
            return "\n".join(text_parts)
        except Exception as e:
            app_logger.error(f"❌ PDF 解析失败: {str(e)}")
            raise
    
    def _parse_docx(self, content: bytes) -> str:
        """
        解析 Word 文档内容
        
        使用 python-docx 提取 Word 文档中的文本内容，包括段落和表格。
        
        Args:
            content: Word 文件的二进制内容
            
        Returns:
            str: 提取的文本内容
            
        Raises:
            Exception: 当 Word 文档解析失败时抛出异常
        """
        from io import BytesIO
        
        text_parts = []
        
        try:
            # 使用 python-docx 打开 Word 文档
            doc = Document(BytesIO(content))
            
            # 遍历每个段落
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:
                    text_parts.append(para_text)
            
            # 遍历表格（如果有）
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            text_parts.append(cell_text)
            
            return "\n".join(text_parts)
        except Exception as e:
            app_logger.error(f"❌ Word 文档解析失败: {str(e)}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """
        清洗文本内容
        
        去除多余的空白符、统一换行符、压缩连续空格等。
        
        Args:
            text: 原始文本
            
        Returns:
            str: 清洗后的文本
        """
        if not text:
            return ""
        
        # 1. 统一换行符为 \n
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        
        # 2. 去除每行首尾的空白
        lines = [line.strip() for line in text.split("\n")]
        
        # 3. 去除空行
        lines = [line for line in lines if line]
        
        # 4. 合并成段落（保留有意义的换行）
        cleaned_text = "\n".join(lines)
        
        # 5. 压缩连续的空格为单个空格
        cleaned_text = re.sub(r" +", " ", cleaned_text)
        
        # 6. 可选：去除中文字符间的多余空格
        # cleaned_text = re.sub(r"(?<=[\u4e00-\u9fff]) (?=[\u4e00-\u9fff])", "", cleaned_text)
        
        return cleaned_text.strip()
